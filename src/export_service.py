from __future__ import annotations

import csv
import io
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

from .constants import DISCLAIMER, EXPORT_DIR
from .utils import safe_timestamp, slugify


class ExportService:
    """Export the reviewed analysis to common BA deliverables."""

    @staticmethod
    def safe_base_name(cr_id: str, title: str) -> str:
        return f"{slugify(cr_id)}_{slugify(title)[:48]}_{safe_timestamp()}"

    @staticmethod
    def _add_table(doc: Document, items: list[dict[str, Any]], columns: list[tuple[str, str]]) -> None:
        if not items:
            doc.add_paragraph("Không có dữ liệu.")
            return
        table = doc.add_table(rows=1, cols=len(columns)); table.style = "Table Grid"
        for index, (_, label) in enumerate(columns):
            table.rows[0].cells[index].text = label
        header_pr = table.rows[0]._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader"); repeat.set(qn("w:val"), "true"); header_pr.append(repeat)
        for item in items:
            cells = table.add_row().cells
            row_pr = table.rows[-1]._tr.get_or_add_trPr(); row_pr.append(OxmlElement("w:cantSplit"))
            for index, (key, _) in enumerate(columns):
                value = item.get(key, "")
                cells[index].text = ", ".join(map(str, value)) if isinstance(value, list) else str(value)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)

    def export_docx(self, cr: dict[str, Any], result: dict[str, Any], output_dir: Path = EXPORT_DIR) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        path = output_dir / f"Impact_Analysis_{self.safe_base_name(cr['id'], cr['title'])}.docx"
        doc = Document()
        doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(10)
        title = doc.add_heading("BÁO CÁO PHÂN TÍCH TÁC ĐỘNG", level=0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph("SmartFactory Change Impact Analyzer"); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("1. Thông tin Change Request", level=1)
        table = doc.add_table(rows=0, cols=2); table.style = "Table Grid"
        for label, key in [("CR ID","id"),("Tiêu đề","title"),("Người yêu cầu","requester"),("Ngày yêu cầu","request_date"),("Mức ưu tiên","priority"),("Loại thay đổi","change_type"),("Mô tả","description"),("Lý do thay đổi","reason_for_change"),("Hành vi mong muốn","expected_behavior")]:
            cells=table.add_row().cells;cells[0].text=label;cells[1].text=str(cr.get(key,""))

        doc.add_heading("2. Tóm tắt phân tích", level=1)
        doc.add_paragraph(result.get("summary", ""))
        doc.add_paragraph(f"Thời gian xử lý: {result.get('processing_time_ms',0)} ms")

        doc.add_heading("3. Tài liệu liên quan", level=1)
        self._add_table(doc,[x for x in result.get("retrieved_artefacts",[]) if x.get("selected",True)],[('rank','Thứ tự'),('document_id','ID'),('artefact_type','Loại'),('title','Tên tài liệu'),('bm25_score','Điểm phù hợp'),('retrieval_reason','Lý do')])

        doc.add_heading("4. Phạm vi tác động", level=1)
        self._add_table(doc,result.get("impacted_modules",[]),[('module_id','Module ID'),('module_name','Module'),('impact_level','Mức tác động'),('impact_reason','Lý do'),('evidence','Căn cứ'),('ba_note','Ghi chú BA')])

        doc.add_heading("5. Artefact cần rà soát", level=1)
        self._add_table(doc,[x for x in result.get("artefacts_to_review",[]) if x.get("selected",True)],[('document_id','ID'),('artefact_type','Loại'),('title','Tên'),('review_action','Hành động'),('ba_note','Ghi chú BA')])

        doc.add_heading("6. Câu hỏi cần làm rõ và rủi ro", level=1)
        for item in result.get("clarifying_questions",[])[:5]: doc.add_paragraph(item.get("question",""),style="List Number")
        for item in result.get("risks",[])[:3]: doc.add_paragraph(f"[{item.get('risk_level','')}] {item.get('risk_description','')} — {item.get('mitigation_suggestion','')}",style="List Bullet")

        doc.add_heading("7. Yêu cầu dự thảo", level=1)
        for item in result.get("draft_user_stories",[])[:2]:
            doc.add_heading(item.get("story_id","User Story"),level=2);doc.add_paragraph(item.get("user_story",""))
            for ac in item.get("acceptance_criteria",[]): doc.add_paragraph(ac,style="List Bullet")
        for item in result.get("draft_business_rules",[])[:2]: doc.add_paragraph(item.get("business_rule",""),style="List Bullet")
        for item in result.get("draft_test_scenarios",[])[:5]: doc.add_paragraph(f"{item.get('test_id','')}: {item.get('test_steps','')} → {item.get('expected_result','')}",style="List Bullet")

        doc.add_heading("8. BA Review", level=1)
        doc.add_paragraph(f"Trạng thái: {result.get('ba_review_status','Draft')}")
        doc.add_paragraph("Thông tin reviewer và nhận xét được lưu trong lịch sử của ứng dụng.")
        doc.save(path)
        return path

    def export_rtm_xlsx(self, cr: dict[str, Any], result: dict[str, Any], output_dir: Path = EXPORT_DIR) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        path = output_dir / f"RTM_{self.safe_base_name(cr['id'], cr['title'])}.xlsx"
        wb=Workbook();ws=wb.active;ws.title="RTM"
        headers=["CR ID","Rule ID","Module ID","Artefact Type","Artefact ID","Test Case ID","Impact Level","Review Status","BA Note"]
        ws.append(headers)
        for row in result.get("traceability_matrix",[]):
            ws.append([row.get("cr_id"),row.get("rule_id"),row.get("module_id"),row.get("artefact_type"),row.get("artefact_id"),row.get("test_case_id"),row.get("impact_level"),row.get("review_status"),row.get("ba_note")])
        fill=PatternFill("solid",fgColor="00327D")
        for cell in ws[1]: cell.fill=fill;cell.font=Font(color="FFFFFF",bold=True);cell.alignment=Alignment(horizontal="center")
        widths=[14,18,18,18,22,20,15,18,35]
        for idx,width in enumerate(widths,1): ws.column_dimensions[chr(64+idx)].width=width
        ws.freeze_panes="A2";ws.auto_filter.ref=ws.dimensions
        wb.save(path);return path

    def export_json(self, cr: dict[str, Any], result: dict[str, Any], output_dir: Path = EXPORT_DIR) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        path=output_dir/f"Analysis_{self.safe_base_name(cr['id'],cr['title'])}.json"
        path.write_text(json.dumps({"change_request":cr,"analysis":result},ensure_ascii=False,indent=2),encoding="utf-8");return path

    def retrieved_csv_bytes(self, result: dict[str, Any]) -> bytes:
        buffer=io.StringIO();headers=["rank","document_id","artefact_type","title","bm25_score","matched_keywords","module_ids","retrieval_reason"]
        writer=csv.DictWriter(buffer,fieldnames=headers);writer.writeheader()
        for item in result.get("retrieved_artefacts",[]):
            writer.writerow({key:", ".join(map(str,item.get(key,[]))) if isinstance(item.get(key),list) else item.get(key,"") for key in headers})
        return buffer.getvalue().encode("utf-8-sig")
